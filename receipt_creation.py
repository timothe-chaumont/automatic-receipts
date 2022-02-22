from typing import List, Dict
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import locale
import datetime as dt
import os
import win32com.client
from dotenv import load_dotenv

from receipt_utils import *

locale.setlocale(locale.LC_ALL, 'fr_FR')

# loads environment variables from .env file
load_dotenv(encoding='utf8')


# Get env constants
CSDESIGN_TRESURER = os.getenv('CSD_TRESURER_NAME')
VIAREZO_TRESURER = os.getenv('VR_TRESURER_NAME')
VR_INFO = os.getenv('VR_INFO')
VR_OFFICIAL_NAME = os.getenv('VR_OFFICIAL_NAME')
VR_IBAN = os.getenv('VR_IBAN')
VR_ACCOUNT_NUMBER = os.getenv('VR_ACCOUNT_NUMBER')

# style constants
BLACK_COLOR = RGBColor(0, 0, 0)
DARK_GREY = RGBColor(100, 100, 100)
FONT_FAMILY = "Arial"


def build_receipt_path(receipt_path: str, month_dir_name: str, receipt_number: str):
    """Builds the path (name) to the receipt file """
    return os.path.join(os.path.join(receipt_path, month_dir_name), receipt_number)


def set_document_styles(styles):
    """Sets the styles of the element of the document (ex: header 1) """

    # base style
    base_style = styles['Normal']
    set_style(base_style, FONT_FAMILY, Pt(9), BLACK_COLOR)
    # custom Title
    new_title_style = styles.add_style('New Title', WD_STYLE_TYPE.PARAGRAPH)
    set_style(new_title_style, FONT_FAMILY, Pt(24), BLACK_COLOR, bold=True)
    # custom heading
    new_heading_style = styles.add_style(
        'New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    set_style(new_heading_style, FONT_FAMILY, Pt(11), BLACK_COLOR, bold=True)
    # custom footer
    new_footer_style = styles.add_style('New Footer', WD_STYLE_TYPE.PARAGRAPH)
    set_style(new_footer_style, FONT_FAMILY, Pt(8), DARK_GREY)


def add_header_section(document: Document):
    """Writes the title, Adds the logo and CS Design's address"""
    # Add the title
    document.add_paragraph('FACTURE', style="New Title")

    # Add the Header (= table containing the logo and CSD's adress)
    table = document.add_table(1, 2)
    par = table.cell(0, 0).paragraphs[0]
    run = par.add_run()
    run.add_picture('logo.png', width=Cm(3))
    p = table.cell(0, 1).paragraphs[0]
    p.add_run(VR_OFFICIAL_NAME).bold = True
    p.add_run("\n" + VR_INFO)
    table.cell(0, 1).paragraphs[0].alignment = 2


def add_details_section(receipt_number, document):
    """Adds a section with the receipt number, current and due date, and payment methods """
    today = dt.date.today()
    due_date = today + dt.timedelta(weeks=2)

    document.add_paragraph('Détails', style="New Heading")
    p = document.add_paragraph()
    p.add_run(f"Numéro de facture ...............{receipt_number}\n")
    p.add_run(
        f"Date de la facture ................{today.strftime('%#d %b %Y')}\n")
    p.add_run("Mode de règlement .............chèque ou virement\n")
    p.add_run(
        f"Date d'échéance .................{due_date.strftime('%#d %b %Y')}")


def add_footer_section(document):
    """Adds the foot notes : whom to address bank notes to ... """
    p = document.add_paragraph(
        "Établir tous les chèques à l'ordre de ARCS - CS Design")
    p.add_run(
        f"\n\nDomiciliation bancaire pour les règlements par virement : {VR_ACCOUNT_NUMBER}")
    p.add_run(f"\n\nIBAN : {VR_IBAN}")
    p.alignment = 1
    p.style = document.styles['New Footer']


def add_orders_table(orders, total_price: str, document):
    """Adds a table filled with the order's details"""
    # create the table
    table = document.add_table(rows=1, cols=5)

    # fill out the columns names
    hdr_cells = table.rows[0].cells
    table_cols = table.columns
    hdr_cells[0].text = 'Qté'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = 'Désignation'
    hdr_cells[2].text = 'Prix unitaire'
    hdr_cells[3].text = 'TVA'
    hdr_cells[4].text = 'Tot. ligne'

    for order in orders:
        row_cells = table.add_row().cells
        row_cells[0].text = str(order["quantity"])
        row_cells[0].paragraphs[0].alignment = 2
        row_cells[1].text = order["designation"]
        row_cells[2].text = order["unit price"]
        row_cells[3].text = "Non applicable"
        row_cells[4].text = order["line total price"]
    # add blank lines under the real orders
    table.add_row()
    table.add_row()
    # add the last row = sum of prices
    row_cells = table.add_row().cells
    row_cells[3].text = "Net à payer"
    row_cells[4].text = total_price
    # make first and last rows' text bold
    make_rows_bold(table.rows[0])
    make_rows_bold(table.rows[-1])
    # set columns width
    table_cols = table.columns
    set_col_width(table_cols[0], Cm(0.5))
    set_col_width(table_cols[1], Cm(10))
    set_col_width(table_cols[2], Cm(3))
    set_col_width(table_cols[3], Cm(3))
    set_col_width(table_cols[4], Cm(3))


def create_receipt_docx(recipient_info: str, orders: List[Dict], receipt_nb: str, total_price: str, file_name: str):
    """Defines and saves a .docx file 

        Args:
            file_name (str): the path and name of the receipt to be created
    """

    # Create a new document
    document = Document()
    styles = document.styles

    set_document_styles(styles)

    add_header_section(document)

    # Address of recipient
    document.add_paragraph('Facturé à', style="New Heading")
    document.add_paragraph(recipient_info)

    # Receipt details
    add_details_section(receipt_nb, document)

    # responsables
    document.add_paragraph('Responsables', style="New Heading")
    p = document.add_paragraph()
    p.add_run(f"{VIAREZO_TRESURER} en qualité de Trésorier de l'ARCS\n")
    p.add_run(f"{CSDESIGN_TRESURER} en qualité de Trésorier du club CS Design")

    # Table of orders
    document.add_paragraph()
    add_orders_table(orders, total_price, document)

    # Foot notes
    document.add_paragraph()
    document.add_paragraph()
    add_footer_section(document)

    document.save(file_name)


def export_receipt_to_pdf(docx_file_name, pdf_file_name):
    """Exports the receipt to a PDF file"""
    wdFormatPDF = 17

    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(docx_file_name)
    doc.SaveAs(pdf_file_name, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()